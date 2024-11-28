from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

def html_to_word(html_content, output_path):
    # Parse the HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create a Word document
    doc = Document()

    # Add a title
    doc.add_heading('Exported HTML Page', level=1)

    # Process sidebar links (if present)
    sidebar = soup.find('div', class_='sidebar')
    if sidebar:
        doc.add_heading('Sidebar Links', level=2)
        for link in sidebar.find_all('a'):
            link_text = link.text
            link_url = link['href']
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(link_text)
            run.font.size = Pt(12)
            run.font.underline = True
            paragraph.add_run(f' ({link_url})')

    # Process main content
    main_content = soup.find('div', class_='content')
    if main_content:
        doc.add_heading('Main Content', level=2)
        for p in main_content.find_all('p'):
            doc.add_paragraph(p.text)

    # Save the document
    doc.save(output_path)

# Example HTML content
html_content = """
<!DOCTYPE html>
<html>
<head>
    <title>Sample HTML</title>
</head>
<body>
    <div class="sidebar">
        <a href="https://example.com/link1">Link 1</a><br>
        <a href="https://example.com/link2">Link 2</a><br>
    </div>
    <div class="content">
        <p>This is a sample paragraph in the main content.</p>
        <p>Another paragraph with <a href="https://example.com">a link</a>.</p>
    </div>
</body>
</html>
"""

# Output file path
output_path = "exported_page.docx"

# Convert the HTML to Word
html_to_word(html_content, output_path)

print(f"Word document saved to {output_path}")
