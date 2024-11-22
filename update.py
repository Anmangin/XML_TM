
import xml.etree.ElementTree as ET
import re
import sys     
import os
import json
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Assurez-vous d'ajouter cette importation
from docx.enum.text import WD_BREAK  # Importation pour les sauts de page

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def remove_details_tags(text):
    # Utilisation de l'expression régulière pour supprimer les balises <details>...</details>
    cleaned_text = re.sub(r'<(th|td)\s+class=["\']check["\'][^>]*>.*?<!--\$htmlbalise-->.', '', text, flags=re.DOTALL)
    # cleaned_text = re.sub(r'<(th|td)\s+class=["\']check["\'][^>]*>.*?<!--\$htmlbalise-->.*?</\1>', '', text, flags=re.DOTALL)


    return cleaned_text

def Get_objt(root, racine, keyname, fields):
        ProForm = {}
        for FWAW in root.iter(racine):
            key = FWAW.findtext(keyname)
            if key:  # Vérifie que la clé existe
                ProForm[key] = {field: FWAW.findtext(field) for field in fields}
        return ProForm

def filtrer_par_cle(dictionnaire, champ, valeur):
    return {key: val for key, val in dictionnaire.items() if val.get(champ) == valeur}





def lire_et_trier_donnees(pathfileXML, config_path='config.json'):
    # Charger la configuration
    with open(config_path, 'r') as f:
        config = json.load(f)

    # Charger et analyser l'XML
    tree = ET.parse(pathfileXML)
    root = tree.getroot()
    version = root.attrib.get('ver')
    data = {"version": version}

    # Extraire et transformer les données à partir du fichier XML selon la config
    for key, params in config.items():
        racine = params["racine"]
        keyname = params["keyname"]
        fields = params["fields"]
        data[key] = Get_objt(root, racine, keyname, fields)

    # Trier les clés spécifiques
    trier_cles = ["ProPatientVisit", "ProVisitForm", "ProFormGroup", "ProGroupItem", "ProCodeListItem"]
    for key in trier_cles:
        if key in data:
            data[key] = dict(sorted(data[key].items(), key=lambda item: int(item[1].get("OrderNo", 0))))

    # Ajouter le champ "Display" pour ProCodeList
    data["ProCodeList"] = ajouter_display_pro_codelist(data)

    # Ajouter le champ "Display" pour ProItem
    data["ProItem"] = ajouter_display_pro_item(data)

    return data


def ajouter_display_pro_codelist(data):
    """Ajoute un champ 'Display' pour les ProCodeList."""
    for key, value in data["ProCodeList"].items():
        # Filtrer les items associés à la ProCodeList actuelle
        filtered_items = {k: v for k, v in data["ProCodeListItem"].items() if v["ProCodeListGuid"] == key}

        # Construire la chaîne de caractères pour le champ Display
        if len(filtered_items) < 15:
            rep = "<br>".join(f"🔘 {item['Value']} - <b>{item['Caption']}</b>" for item in filtered_items.values())
        else:
            rep = "🔘 Radio bouton trop long"

        # Ajouter le champ Display
        data["ProCodeList"][key]["Display"] = rep

    return data["ProCodeList"]


def ajouter_display_pro_item(data):
    """Ajoute un champ 'Display' pour les ProItem."""
    for key, value in data["ProItem"].items():
            if value["ProDataTypeId"] == "5":
                rep = "📅 DD/MM/YYYY"
            else:
                rep = f"{value['SasType']} - {value['MaxLength']}"

            # Ajouter le champ Display
            if rep :data["ProItem"][key]["Display"] = rep





            # Partie du la nature de la var
            if value["Hidden"]=="True": Hidden="👻"
            else: Hidden=""
            if value["ReadOnly"]=="True": ReadOnly="🔒"
            else: ReadOnly=""
            value["Status"]= Hidden + ReadOnly


    return data["ProItem"]

# Appel de la fonction

def find_arbo(TBNode):
            if TBNode["ParentTBNodeId"]=="6":source="Fiche"
            elif TBNode["ParentTBNodeId"]=="7":source="groupe"
            elif TBNode["ParentTBNodeId"]=="8":source="item"
            elif TBNode["ParentTBNodeId"]=="9":source="codelist"
            elif int(TBNode["ParentTBNodeId"]) < 10 : source="unknown"
            else : source="unknown"
            return source
    




def get_message(data,ProItemGuid,ProGroupGuid,ProFormGuid):
    """genere le message pour les editcheck"""
    Message=""
    i=0
    ProEdit=filtrer_par_cle( data,"ProItemGuid",ProItemGuid)
    for  key,Edit in ProEdit.items():
        if (Edit["ProItemGuid"] == ProItemGuid) and (Edit["TargetProFormGuid"]==ProFormGuid or Edit["TargetProFormGuid"] is None) and (Edit["TargetProGroupGuid"]==ProGroupGuid or Edit["TargetProGroupGuid"] is None)  :
            i+=1


            EditACtion= Edit["ProEditActionId"]
            if   EditACtion =="1"           :EditACtion="Valid"
            elif EditACtion =="3"           :EditACtion="Enabled"
            elif EditACtion =="6"           :EditACtion="Hidden"
            elif EditACtion =="10"          :EditACtion="DVA"
            elif EditACtion =="11"          :EditACtion="Email"
            elif EditACtion =="9"           :EditACtion="DVC"
            elif EditACtion =="23"          :EditACtion="Dynamic codelist filter"
            msg=Edit["Message"]
            chk=Edit["ActionExpression"]
            DTE=Edit['DataExpression']
            path=Edit["TargetPath"]

            Message+=f"<tr><td> {EditACtion}:{path}</td> </tr>"
            Message+=f"<tr> <td> <pre><code class='javascript'>#Action Expression \n{chk} \n#data Expression \n{DTE} \n</code></pre> </td>"
            Message+=f"<td> {msg}</td> </tr>"
            
            # if EditACtion=="Valid": 
            #     check={ 
            #         "ActionExpression":chk   ,
            #         "Message":msg,
            #         "path":path ,
            #         'i_SasName':i_SasName,
            #         'F_SasName':F_SasName                               
            #         }
            #     checks.append(check)





    Message+="</table></details>"
    if i==0:Message=""
    else: Message= f"<details> <summary>{i} EditCheck </summary><table>" + Message


    return Message




def exporter_donnees_markdown_eCRF(data,ACTversion,display_Edit=True):

        
 
        content=f"<body>\n\n\n"

        
        if ACTversion != version :
             content+=f"# Fichier non compatible avec la version en cours de TM : actuel:{OngoingVer}, version du fichier : {version}  \n"
             
        


        # errortitle=0      
        # for  PVkey,TBNode in data["TBNode"].items():    
        #     if TBNode["TBNodeId"] is not None  :
        #         TBNodeId=TBNode["TBNodeId"]
        #         Caption=TBNode["Caption"]
        #         listItm=filtrer_par_cle( data["TBNode"],"ParentTBNodeId",TBNodeId)
        #         if len(listItm)==0 :
        #             if errortitle==0:
        #                 errortitle+=1
        #                 content+=f"# ERREUR DANS LES DOSSIERS DU FICHIER  \n"
        #                 content+=f"{find_arbo(TBNode)}/{Caption} est vide TBNodeId:{TBNodeId}\n"

        # Parcourir les données
        if logprint:print(len(data["ProPatientVisit"]), "Patients visit")
        if len(data["ProPatientVisit"])==0 : 
            if errortitle==0:
                        errortitle+=1
                        content+=f"# ERREUR DANS LES DOSSIERS DU FICHIER  \n"
            content+=f"Pas d'arborescence patient ! \n"
        for  PVkey,PatientVisit in data["ProPatientVisit"].items():



            # Ecriture des variables nécéssaire pour l'affichage
            ProVisitGuid        =PatientVisit["ProVisitGuid"]
            V_description       =data["ProVisit"][ProVisitGuid]["Description"]
            V_Caption           =data["ProVisit"][ProVisitGuid]["Caption"]
            V_OrderNo           =PatientVisit["OrderNo"]



            if not unic_form: # si on fait un CRF complet, pas la peine de rajouter la liste des visites pour chaque fiche.
                content+=f"<h1> {V_description} </h1> \n"
                if logprint:("écriture de #",V_description)



            ProVisitForm=filtrer_par_cle( data["ProVisitForm"],"ProVisitGuid",ProVisitGuid)
           
            for  VFkey,VisitForm in ProVisitForm.items():
                ProFormGuid=VisitForm.get("ProFormGuid")
                written=data["ProForm"][ProFormGuid].get("written", False)
                if written and unic_form: continue    

                # Ecriture des variables nécéssaire pour l'affichage
                F_OrderNo           =VisitForm["OrderNo"]
                F_description       =data["ProForm"][ProFormGuid]["Description"]
                F_SasName           =data["ProForm"][ProFormGuid]["SasName"]
                F_Caption           =data["ProForm"][ProFormGuid]["Caption"]


                data["ProForm"][ProFormGuid]["written"]=True
                content+=f"<h2> {F_description}  </h2>\n"
                if logprint:print("écriture de ##",F_description)

                if unic_form:# cas ou on imprime qu'une version de la fiche, alors on liste les visites ou elle apparait.
                    ListVisit=filtrer_par_cle( data["ProVisitForm"],"ProFormGuid",ProFormGuid)
                    ListVisit= dict(sorted(ListVisit.items(), key=lambda item: int(item[1].get("OrderNo", 0))))
                    Li="Liste des visites avec cette fiches :"
                    for  VFkey,LV in ListVisit.items():
                        # Ecriture des variables nécéssaire pour l'affichage
                        Vgui_temp=LV["ProVisitGuid"]
                        V_description_temp = data["ProVisit"][Vgui_temp]["Description"]
                        Li+= f"{V_description_temp}"
                    content+=f"<p>{Li}</p> \n"

                    



                ProFormGroup=filtrer_par_cle( data["ProFormGroup"],"ProFormGuid",ProFormGuid)
                for  key,FormGroup in ProFormGroup.items():
                    # Ecriture des variables nécéssaire pour l'affichage
                    ProGroupGuid        =FormGroup["ProGroupGuid"]
                    G_description       =data["ProGroup"][ProGroupGuid]["Caption"]
                    G_Caption           =data["ProGroup"][ProGroupGuid]["Caption"]


                    # Titre du groupe, ainsi que l'entete pour les  variables
                    content+=f"<h3> {G_description} </h3> \n"
                                    
                    content+="<table style='width:100%;'>\n"
                    content+="<tr>\n"
                    content+="<th style='width:600px; text-align:center;'><strong>Label de la question </strong></th>\n"
                    if display_Edit:content+="<th class='check' style='width:300px; text-align:center;'><strong>Check</strong></th> <!--$htmlbalise-->\n"
                    content+="<th style='width:300px; text-align:center;'><strong>Réponses possibles</strong></th>\n"
                    content+="<th style='width:50px; text-align:center;'><strong>Sas</strong></th>\n"
                    content+="</tr>\n"



                    ProGroupItem=filtrer_par_cle( data["ProGroupItem"],"ProGroupGuid",ProGroupGuid)
                    if logprint:print("écriture de ###",G_description, "(", len(ProGroupItem),")")
                    for  GI_key,GroupItem in ProGroupItem.items():
                        # Ecriture des variables nécéssaire pour l'affichage
                        I_OrderNo           =GroupItem["OrderNo"]
                        ProItemGuid         =GroupItem.get("ProItemGuid")
                        I_description       =data["ProItem"][ProItemGuid]["Description"]
                        i_SasName           =data["ProItem"][ProItemGuid]["SasName"]
                        I_Caption           =data["ProItem"][ProItemGuid]["Caption"]
                        ProCodeListGuid     =data["ProItem"][ProItemGuid]["ProCodeListGuid"]
                        Hidden              =data["ProItem"][ProItemGuid]["Hidden"]
                        Disabled            =data["ProItem"][ProItemGuid]["Disabled"]
                        ReadOnly            =data["ProItem"][ProItemGuid]["ReadOnly"]
                        i_Display           =data["ProItem"][ProItemGuid]["Display"]
                        I_Status            =data["ProItem"][ProItemGuid]["Status"]
                        if ProCodeListGuid :rep=data["ProCodeList"][ProCodeListGuid]["Display"]
                      
                      
                        # DICO[f"[{V_Caption}][{F_Caption}][{G_Caption}][{I_Caption}] "] = {

                        #     'i_SasName':i_SasName,
                        #     'F_SasName':F_SasName     
                        # }

                        
                        Message = get_message(data["ProEdit"],ProItemGuid,ProGroupGuid,ProFormGuid)    
                        
                        content+=" <tr> \n"                       
                        content+=            f" <td style='width:600px; text-align:left;'> {I_description}</td>\n" 
                        if display_Edit:content+=             f" <td class='check' style='width:600px; text-align:left;'>  {Message} </td>\n" 
                        content+=             f" <td style='width:300px; text-align:center;'> {rep} </td> \n"
                        content+=            f"<td style='width:50px; text-align:center; color:red; font-size: 10px;'> <b> {I_Status}{i_SasName} </b></td> \n" 
                        content+=" </tr>\n"            
                        # print("écriture de ",I_description)
                        

                    content+="</table>\n\n"
                content+=f"<span style='font-size:11.0pt;line-height:107%;font-family:""Calibri"",sans-serif;mso-ascii-theme-font:minor-latin;mso-fareast-font-family:Calibri;mso-fareast-theme-font:minor-latin;mso-hansi-theme-font:minor-latin;mso-bidi-font-family:""Times New Roman"";mso-bidi-theme-font:minor-bidi;mso-ansi-language:FR;mso-fareast-language:EN-US;mso-bidi-language:AR-SA'><br clear=all style='mso-special-character:line-break;page-break-before:always'></span>"
        content+=f"</body>\n\n\n"
        return content
        




        if logprint:print(f"fin d'écriture dans {pathout}")


logprint=False
unic_form=True

DICO={}
checks=[]
Pathin = sys.argv[1]
file_t = os.path.basename(Pathin)
file=file_t.replace('.xml', '')


directory = os.path.dirname(Pathin)
path= os.path.dirname(os.path.dirname(directory)) + r"\OUTFILE"


config_path=os.path.dirname(os.path.dirname(directory)) + "/config.json"
data = lire_et_trier_donnees(Pathin,config_path)

version=data["version"]
        #content+=f"# Version de TB pour ce fichier : {version}  \n")

with open(f"{os.path.dirname(directory)}/VersionTM.txt", "r", encoding="utf-8") as fichier:
    ACTversion = fichier.read()


# content = exporter_donnees_markdown_eCRF(data,ACTversion,False)

def html_to_docx(html_content, output_path):
    """
    Convertit le contenu HTML en un document Word (.docx), incluant les tableaux et les sauts de page.
    
    :param html_content: Le contenu HTML à convertir.
    :param output_path: Le chemin où enregistrer le fichier .docx.
    """
    # Parse le HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Crée un document Word
    doc = Document()
    
    # Boucle à travers les éléments HTML
    for element in soup.body.descendants:
        if element.name == 'p':
            # Ajout d'un paragraphe normal
            para = doc.add_paragraph(element.get_text())
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif element.name == 'h1':
            # Ajout d'un titre de niveau 1
            para = doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            # Ajout d'un titre de niveau 2
            para = doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            # Ajout d'un titre de niveau 3
            para = doc.add_heading(element.get_text(), level=3)
        elif element.name == 'ul':
            # Liste à puces
            for li in element.find_all('li'):
                para = doc.add_paragraph(f"- {li.get_text()}", style='List Bullet')
        elif element.name == 'ol':
            # Liste numérotée
            for li in element.find_all('li'):
                para = doc.add_paragraph(f"{li.get_text()}", style='List Number')
        elif element.name == 'table':
            # Gestion des tableaux
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text(strip=True)
        elif element.name == 'span' and 'page-break-before' in element.get('style', ''):
            # Saut de page détecté par la balise <span> avec un style page-break-before:always
            doc.add_paragraph().clear()  # Ajoute un paragraphe vide avant le saut de page
            doc.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)  # Ajoute le saut de page
    
    # Enregistre le document
    doc.save(output_path)
    print(f"Document enregistré sous : {output_path}")



def exporter_donnees_word_eCRF(data, ACTversion,filepath, display_Edit=True, logprint=True,):
    doc = Document()
    
    # Start by checking the version
    if ACTversion != version:
        doc.add_heading(f"Fichier non compatible avec la version en cours de TM : actuel:{OngoingVer}, version du fichier : {version}", level=1)
    
    # Print error if no patient visits are found
    if len(data["ProPatientVisit"]) == 0:
        doc.add_heading("ERREUR DANS LES DOSSIERS DU FICHIER", level=1)
        doc.add_paragraph("Pas d'arborescence patient !")

    # Iterate through patient visits
    for PVkey, PatientVisit in data["ProPatientVisit"].items():
        ProVisitGuid = PatientVisit["ProVisitGuid"]
        V_description = data["ProVisit"][ProVisitGuid]["Description"]
        V_Caption = data["ProVisit"][ProVisitGuid]["Caption"]
        V_OrderNo = PatientVisit["OrderNo"]

        # Add patient visit heading
        doc.add_heading(f"{V_description}", level=2)

        ProVisitForm = filtrer_par_cle(data["ProVisitForm"], "ProVisitGuid", ProVisitGuid)
        for VFkey, VisitForm in ProVisitForm.items():
            ProFormGuid = VisitForm.get("ProFormGuid")
            written = data["ProForm"][ProFormGuid].get("written", False)
            if written and unic_form:
                continue

            # Write form description
            F_description = data["ProForm"][ProFormGuid]["Description"]
            F_SasName = data["ProForm"][ProFormGuid]["SasName"]
            F_Caption = data["ProForm"][ProFormGuid]["Caption"]

            # Mark the form as written
            data["ProForm"][ProFormGuid]["written"] = True
            doc.add_heading(f"{F_description}", level=3)

            # Write visit listing if unic_form
            if unic_form:
                ListVisit = filtrer_par_cle(data["ProVisitForm"], "ProFormGuid", ProFormGuid)
                ListVisit = dict(sorted(ListVisit.items(), key=lambda item: int(item[1].get("OrderNo", 0))))
                Li = "Liste des visites avec cette fiche :"
                for VFkey, LV in ListVisit.items():
                    Vgui_temp = LV["ProVisitGuid"]
                    V_description_temp = data["ProVisit"][Vgui_temp]["Description"]
                    Li += f"{V_description_temp}"
                doc.add_paragraph(Li)

            # Write group items
            ProFormGroup = filtrer_par_cle(data["ProFormGroup"], "ProFormGuid", ProFormGuid)
            for key, FormGroup in ProFormGroup.items():
                ProGroupGuid = FormGroup["ProGroupGuid"]
                G_description = data["ProGroup"][ProGroupGuid]["Caption"]
                doc.add_heading(f"{G_description}", level=4)

                # Add table for group items
                table = doc.add_table(rows=1, cols=4)
                table.autofit = True
                table.style = 'Table Grid'

                # Table header
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Label de la question'
                if display_Edit:
                    hdr_cells[1].text = 'Check'
                hdr_cells[2].text = 'Réponses possibles'
                hdr_cells[3].text = 'Sas'

                ProGroupItem = filtrer_par_cle(data["ProGroupItem"], "ProGroupGuid", ProGroupGuid)
                for GI_key, GroupItem in ProGroupItem.items():
                    I_description = data["ProItem"][GroupItem["ProItemGuid"]]["Description"]
                    I_Status = data["ProItem"][GroupItem["ProItemGuid"]]["Status"]
                    i_SasName = data["ProItem"][GroupItem["ProItemGuid"]]["SasName"]
                    Message = get_message(data["ProEdit"], GroupItem["ProItemGuid"], ProGroupGuid, ProFormGuid)
                    ProCodeListGuid     =data["ProItem"][GroupItem["ProItemGuid"]]["ProCodeListGuid"]
                    if ProCodeListGuid :rep=data["ProCodeList"][ProCodeListGuid]["Display"]

                    para = doc.add_paragraph()
                    run = para.add_run("Voici un emoji : 🔘")

                    # Add row for each group item
                    row_cells = table.add_row().cells
                    row_cells[0].text = I_description
                    row_cells[2].text = rep
                    p = row_cells[2].paragraphs[0]  # Premier paragraphe de la cellule
                    run = p.add_run(rep)

                    # Appliquer la police compatible avec les émojis
                    run.font.name = 'Segoe UI Emoji'

                    # Modifier également au niveau XML pour garantir la compatibilité dans Word
                    r = run._element
                    rPr = r.get_or_add_rPr()
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), 'Segoe UI Emoji')
                    rFonts.set(qn('w:hAnsi'), 'Segoe UI Emoji')
                    rPr.append(rFonts)
                    row_cells[3].text = f"{I_Status} {i_SasName}"

            doc.add_page_break()

    # Save the document


    doc.save(filepath)





    if logprint:
        print("Document saved as 'exported_data.docx'")

# Example usage

content=""
exporter_donnees_word_eCRF(data,ACTversion,f"{path}/HTML/{file}.docx",False)


# html_to_docx(htmlcontent,f"{path}/HTML/{file}.docx")

with open( f"{path}/HTML/{file}.html" , 'w', encoding='utf-8') as f:
            htmlcontent= f"    <head>   \n    <meta charset='UTF-8'>    \n   <meta name='viewport' content='width=device-width, initial-scale=1.0'>\n <title>test fiche</title> \n   <link rel='stylesheet' href='style.css'> \n </head> \n\n" +  content #remove_details_tags(content)
            f.write(htmlcontent)

with open( f"{path}/MD/{file}.md" , 'w', encoding='utf-8') as f:
            f.write(content)
        # json_output = json.dumps(checks, indent=4)
        # with open( f"{pathout}/JSON/{file}.json" , 'w', encoding='utf-8') as f:
        #      f.write(json_output)
        # json_DICO = json.dumps(DICO, indent=4)
        # with open( f"{pathout}/JSON/DICO.json" , 'w', encoding='utf-8') as f:
        #      f.write(json_DICO)
        # JSON_FUNC = json.dumps(data["ProScriptFunction"], indent=4)
        # with open( f"{pathout}/JSON/Function.json" , 'w', encoding='utf-8') as f:
        #      f.write(JSON_FUNC)


             
